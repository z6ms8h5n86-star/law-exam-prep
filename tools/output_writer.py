#!/usr/bin/env python3
"""
Multi-format output writer for law-import skill.

Usage:
    python output_writer.py --input notes/ --format md --output output/
    python output_writer.py --input notes/ --format html --template law_note
    python output_writer.py --input notes/ --format pdf
    python output_writer.py --input notes/ --format docx
    python output_writer.py --input notes/ --format all

Generates MD/HTML/PDF/DOCX from the structured course directories.
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path
from datetime import datetime
from typing import Optional

TEMPLATE_DIR = Path(__file__).parent.parent / "templates"

def discover_courses(input_dir: str) -> dict:
    """Discover course directories and their subdirectories."""
    courses = {}
    root = Path(input_dir)
    if not root.is_dir():
        return courses

    for item in root.iterdir():
        if item.is_dir() and not item.name.startswith("_") and not item.name.startswith("."):
            subdirs = {}
            for sub in item.iterdir():
                if sub.is_dir():
                    md_files = list(sub.glob("*.md"))
                    if md_files:
                        subdirs[sub.name] = [str(f) for f in md_files]
            if subdirs:
                courses[item.name] = {
                    "path": str(item),
                    "subdirs": subdirs,
                }
    return courses

def read_file_safe(filepath: str) -> str:
    """Read file with encoding fallback."""
    for enc in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            with open(filepath, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()

def parse_frontmatter(text: str) -> tuple:
    """Parse YAML frontmatter from Markdown. Returns (frontmatter_dict, body_text)."""
    fm = {}
    body = text

    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3:
            fm_text = parts[1]
            body = parts[2]
            # Simple YAML parsing (key: value)
            for line in fm_text.strip().split("\n"):
                line = line.strip()
                if ":" in line and not line.startswith("#"):
                    key, _, value = line.partition(":")
                    fm[key.strip()] = value.strip().strip('"').strip("'")
    return fm, body.strip()

def generate_md(courses: dict, output_dir: str) -> list:
    """Generate/consolidate Markdown output."""
    written = []
    out_path = Path(output_dir)

    # Generate course overview files
    for course_name, course_info in courses.items():
        course_dir = out_path / course_name
        course_dir.mkdir(parents=True, exist_ok=True)

        # Course MOC
        moc_lines = [
            f"# {course_name}",
            "",
            f"> 由 law-import 自动生成 — {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            "",
            "## 目录",
            "",
        ]
        for sub_name, files in course_info["subdirs"].items():
            moc_lines.append(f"### {sub_name}")
            for f in files:
                rel_path = os.path.relpath(f, str(course_dir))
                title = Path(f).stem
                moc_lines.append(f"- [{title}]({rel_path})")
            moc_lines.append("")

        moc_file = course_dir / f"{course_name}概述.md"
        with open(moc_file, "w", encoding="utf-8") as f:
            f.write("\n".join(moc_lines))
        written.append(str(moc_file))

        # Copy/conversion is handled by format_converter; here we just ensure the MD files exist
        for sub_name, files in course_info["subdirs"].items():
            sub_dir = course_dir / sub_name
            sub_dir.mkdir(parents=True, exist_ok=True)
            for src_file in files:
                dest = sub_dir / Path(src_file).name
                if not dest.exists():
                    content = read_file_safe(src_file)
                    with open(dest, "w", encoding="utf-8") as f:
                        f.write(content)
                    written.append(str(dest))

    # Global index
    index_lines = [
        "# 法学总览",
        "",
        f"> 由 law-import 自动生成 — {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
        "## 课程列表",
        "",
    ]
    for course_name in sorted(courses.keys()):
        index_lines.append(f"- [[{course_name}概述|{course_name}]]")
    index_lines.append("")
    index_lines.append("## 跨课程资源")
    index_lines.append("")
    index_lines.append("- [_跨课程/法条](_跨课程/法条/)")
    index_lines.append("- [_跨课程/术语](_跨课程/术语/)")
    index_lines.append("- [_跨课程/文献](_跨课程/文献/)")
    index_lines.append("- [_原始文件/原始文件索引](_原始文件/原始文件索引.md)")
    index_lines.append("- [_原始文件/项目进度](_原始文件/项目进度.md)")

    index_file = out_path / "法学总览.md"
    with open(index_file, "w", encoding="utf-8") as f:
        f.write("\n".join(index_lines))
    written.append(str(index_file))

    return written

def generate_html(courses: dict, output_dir: str, template_name: str = "law_note") -> list:
    """Generate HTML output from course structure."""
    written = []
    out_path = Path(output_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    # Try to load template
    template_path = TEMPLATE_DIR / f"{template_name}.html"
    template_html = ""
    if template_path.exists():
        template_html = read_file_safe(str(template_path))
    else:
        # Inline minimal template
        template_html = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{{TITLE}}</title>
<style>
body { font-family: -apple-system, "Microsoft YaHei", sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.8; }
h1 { border-bottom: 2px solid #333; padding-bottom: 10px; }
h2 { border-bottom: 1px solid #ccc; padding-bottom: 5px; margin-top: 30px; }
h3 { margin-top: 20px; color: #555; }
table { border-collapse: collapse; width: 100%; margin: 15px 0; }
th, td { border: 1px solid #ddd; padding: 8px 12px; text-align: left; }
th { background: #f5f5f5; }
.exam-critical { color: #d32f2f; font-weight: bold; }
.exam-important { color: #e67e00; font-weight: bold; }
.note { color: #666; font-style: italic; }
blockquote { border-left: 4px solid #2196f3; padding: 10px 15px; margin: 15px 0; background: #f0f7ff; }
@media print { body { font-size: 12pt; } }
</style>
</head>
<body>
{{BODY}}
</body>
</html>"""

    for course_name, course_info in courses.items():
        body_parts = [f"<h1>{course_name}</h1>",
                      f"<p><em>由 law-import 生成 — {datetime.now().strftime('%Y-%m-%d')}</em></p>"]

        for sub_name in ["笔记", "重点与考纲", "试题", "案例", "法条", "术语", "文献"]:
            if sub_name in course_info["subdirs"]:
                body_parts.append(f"<h2>{sub_name}</h2>")
                for filepath in course_info["subdirs"][sub_name]:
                    content = read_file_safe(filepath)
                    fm, body = parse_frontmatter(content)

                    # Convert Markdown to basic HTML
                    html_body = body
                    html_body = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html_body, flags=re.MULTILINE)
                    html_body = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html_body, flags=re.MULTILINE)
                    html_body = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html_body, flags=re.MULTILINE)
                    html_body = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html_body)
                    html_body = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html_body)
                    html_body = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', html_body)
                    html_body = re.sub(r'\n\n', '</p><p>', html_body)
                    html_body = f"<p>{html_body}</p>"
                    html_body = html_body.replace("<p></p>", "")

                    # Convert tables (basic)
                    html_body = re.sub(r'\|(.+)\|\n\|[-| ]+\|\n((?:\|.+\|\n?)*)',
                                       r'<table><tr>\1</tr>\2</table>', html_body)

                    # Highlight exam markers
                    for marker in ["【★重点】", "【必考】", "【历年真题】"]:
                        html_body = html_body.replace(marker, f'<span class="exam-critical">{marker}</span>')
                    for marker in ["【高频】", "【重点】"]:
                        html_body = html_body.replace(marker, f'<span class="exam-important">{marker}</span>')

                    title = fm.get("title", Path(filepath).stem)
                    body_parts.append(f"<h3>{title}</h3>")
                    body_parts.append(html_body)

        html_content = template_html.replace("{{TITLE}}", f"{course_name} — 法学复习资料")
        html_content = html_content.replace("{{BODY}}", "\n".join(body_parts))

        html_file = out_path / f"{course_name}.html"
        with open(html_file, "w", encoding="utf-8") as f:
            f.write(html_content)
        written.append(str(html_file))

    # Index page
    index_body = ["<h1>法学复习资料</h1>", "<ul>"]
    for course_name in sorted(courses.keys()):
        index_body.append(f'<li><a href="{course_name}.html">{course_name}</a></li>')
    index_body.append("</ul>")
    index_html = template_html.replace("{{TITLE}}", "法学复习资料 — 总览")
    index_html = index_html.replace("{{BODY}}", "\n".join(index_body))
    index_file = out_path / "index.html"
    with open(index_file, "w", encoding="utf-8") as f:
        f.write(index_html)
    written.append(str(index_file))

    return written

def generate_pdf(courses: dict, output_dir: str) -> list:
    """Generate PDF output. First generates HTML, then converts via weasyprint if available."""
    written = []
    out_path = Path(output_dir)

    try:
        from weasyprint import HTML
    except ImportError:
        # Fallback: just generate HTML and note that PDF requires weasyprint
        html_files = generate_html(courses, output_dir)
        print("WARNING: weasyprint not installed. Only HTML generated.", file=sys.stderr)
        print("Install: pip install weasyprint", file=sys.stderr)
        return html_files

    # Generate HTML first
    html_files = generate_html(courses, output_dir)
    for hf in html_files:
        pdf_path = hf.replace(".html", ".pdf")
        try:
            HTML(filename=hf).write_pdf(pdf_path)
            written.append(pdf_path)
        except Exception as e:
            print(f"PDF conversion failed for {hf}: {e}", file=sys.stderr)

    return written

def generate_docx(courses: dict, output_dir: str) -> list:
    """Generate DOCX output."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        return []

    written = []
    out_path = Path(output_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    for course_name, course_info in courses.items():
        doc = Document()

        # Title
        title = doc.add_heading(course_name, level=0)
        doc.add_paragraph(f"由 law-import 生成 — {datetime.now().strftime('%Y-%m-%d')}")

        for sub_name in ["笔记", "重点与考纲", "试题", "案例", "法条", "术语", "文献"]:
            if sub_name in course_info["subdirs"]:
                doc.add_heading(sub_name, level=1)
                for filepath in course_info["subdirs"][sub_name]:
                    content = read_file_safe(filepath)
                    fm, body = parse_frontmatter(content)
                    note_title = fm.get("title", Path(filepath).stem)
                    doc.add_heading(note_title, level=2)

                    # Add body text paragraph by paragraph
                    for para_text in body.split("\n\n"):
                        para_text = para_text.strip()
                        if not para_text:
                            continue
                        # Skip if it looks like a heading (will be handled differently)
                        if re.match(r'^#{1,3}\s', para_text):
                            heading_text = re.sub(r'^#+\s+', '', para_text)
                            doc.add_heading(heading_text, level=3)
                        else:
                            p = doc.add_paragraph(para_text)
                            # Highlight exam markers
                            for marker in ["【★重点】", "【必考】", "【历年真题】"]:
                                if marker in para_text:
                                    for run in p.runs:
                                        if marker in run.text:
                                            run.bold = True
                                            run.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)

        docx_file = out_path / f"{course_name}.docx"
        doc.save(str(docx_file))
        written.append(str(docx_file))

    return written

def main():
    parser = argparse.ArgumentParser(description="Multi-format output writer for law-import")
    parser.add_argument("--input", "-i", required=True, help="Input directory with course structure")
    parser.add_argument("--format", "-f", choices=["md", "html", "pdf", "docx", "all"],
                        default="md", help="Output format")
    parser.add_argument("--output", "-o", default="./output", help="Output directory")
    parser.add_argument("--template", default="law_note", help="HTML template name")
    args = parser.parse_args()

    if not os.path.isdir(args.input):
        print(f"ERROR: Input directory not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    courses = discover_courses(args.input)
    if not courses:
        print(f"WARNING: No course directories found in {args.input}", file=sys.stderr)

    all_written = []

    if args.format in ("md", "all"):
        files = generate_md(courses, args.output)
        all_written.extend(files)
        print(f"MD: {len(files)} files → {args.output}", file=sys.stderr)

    if args.format in ("html", "all"):
        files = generate_html(courses, args.output, args.template)
        all_written.extend(files)
        print(f"HTML: {len(files)} files → {args.output}", file=sys.stderr)

    if args.format in ("pdf", "all"):
        files = generate_pdf(courses, args.output)
        all_written.extend(files)
        print(f"PDF: {len(files)} files → {args.output}", file=sys.stderr)

    if args.format in ("docx", "all"):
        files = generate_docx(courses, args.output)
        all_written.extend(files)
        print(f"DOCX: {len(files)} files → {args.output}", file=sys.stderr)

    # Output summary
    summary = {
        "format": args.format,
        "output_dir": os.path.abspath(args.output),
        "files_written": len(all_written),
        "file_list": all_written,
    }
    print(json.dumps(summary, indent=2, ensure_ascii=False))

if __name__ == "__main__":
    main()
